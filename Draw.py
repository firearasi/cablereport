import os
from datetime import date, datetime, time
from babel.dates import format_date, format_datetime, format_time

import dateutil.parser as dateparser
import argparse
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as plticker


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches



#还要分端口改
def read_data_from_file(fname):
    with open(fname, 'r') as file:
        lines = file.readlines()
        lines=[s.strip() for s in lines]
        result = []
        data_lines=list(range(12, len(lines),2))
        time_stamp = lines[4][6:]
        time=dateparser.parse(time_stamp)
        for l in data_lines:
            result.append(lines[l].split(None)[1:])
        return (time, np.array(result, dtype='float16'))
        

parser = argparse.ArgumentParser(description='Data path name.')

#命令行参数
parser.add_argument('path', help='数据文件夹')
parser.add_argument('--product', help='产品名称', default='')
parser.add_argument('--samplenumber', help='样品编号', default='')
parser.add_argument('--institution', help='受检单位', default='')
parser.add_argument('--type', help='型号规格', default='')

#以上是命令行参数，典型的命令行调用是
# python Draw.py 文件夹 
# python Draw.py 文件夹 --product x -- samplenumber y --institution z --type s




args = parser.parse_args()
print(args.path)
path = args.path
product = args.product
samplenumber = args.samplenumber
institution = args.institution
type = args.type


files = []
for file in os.listdir(args.path):
    if file.endswith(".txt"):
        files.append(os.path.join(args.path, file))
files.sort()

# 处理初始文件，获取一些信息
fname0 = files[0]
print(fname0)
with open(fname0, 'r') as file0:
    lines0 = file0.readlines()
    lines0=[s.strip() for s in lines0]
    time_stamp = lines0[4][6:]
    terminal_lines = range(12, len(lines0),2)
    num_terminal= len(terminal_lines)
    
    # 在初始文件中寻找端口号码
    terminal_names = [lines0[i].split()[0] for i in terminal_lines]

base_data=read_data_from_file(fname0)[1]
data=np.zeros(shape=(0, 3))
delta=np.zeros(shape=(0, 3))
datetimes=[]
for fname in files:
    read_data=read_data_from_file(fname)
    new_delta=read_data[1]-base_data
    data = np.vstack((data, read_data[1]))
    delta = np.vstack((delta, new_delta))
    datetimes.append(read_data[0])


   
# 画图
length=int(len(delta)/num_terminal) 

# x是时间戳 小时
x = [(item-datetimes[0]).total_seconds()/3600.0 for item in datetimes ]

from matplotlib.font_manager import FontProperties
font = FontProperties(fname=r"simsun.ttc", size=16)


for i in range(num_terminal):
    plt.figure(1, figsize=(16, 4))
    plt.suptitle(u'端口' + terminal_names[i] + '插入损耗变化量', fontproperties=font)
    
    ax = plt.subplot(131)
    ax.plot(x, delta[i::num_terminal, 0])
    ax.set_title('1310nm')
    ax.set_ylim(-0.5, 0.5)
    ax.set_xlabel(r'时间 (h)', fontproperties=font)
    ax.set_ylabel(r'变化量 (dB)', fontproperties=font)
    
    #设置y轴数字间隔
    loc = plticker.MultipleLocator(base=0.1) 
    ax.yaxis.set_major_locator(loc)
    
    ax = plt.subplot(132)
    ax.plot(x, delta[i::num_terminal, 1])  
    ax.set_title('1490nm')
    ax.set_ylim(-0.5, 0.5)
    ax.set_xlabel(r'时间 (h)', fontproperties=font)
    #ax.set_ylabel(r'变化量 (dB)', fontproperties=font)
    loc = plticker.MultipleLocator(base=0.1) 
    ax.yaxis.set_major_locator(loc)
    
    
    ax = plt.subplot(133)
    ax.plot(x, delta[i::num_terminal, 2])
    ax.set_ylim(-0.5, 0.5)
    ax.set_title('1550nm')
    ax.set_xlabel(r'时间 (h)', fontproperties=font)
    #ax.set_ylabel(r'变化量 (dB)', fontproperties=font)
    loc = plticker.MultipleLocator(base=0.1)
    ax.yaxis.set_major_locator(loc)
    
   
    plt.tight_layout(pad=3, w_pad=4, h_pad=4)
    plt.savefig(str(i) + '.jpg')
    
    #plt.show()
    
#生成报告
    
rounds = num_terminal // 2

for j in range(rounds):
    # 从文件中读取的开始试验时间和结束试验时间
    start_time_string = format_datetime(datetimes[0], locale='zh_CN')
    end_time_string = format_datetime(datetimes[-1], locale='zh_CN')
    
    
    report = Document()
    
    title = report.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设定标题字体和内容
    run = title.add_run()
    run.font.name = '宋体'
    run.font.size = Pt(12)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.add_text('试验名称：XXX试验插入损耗变化量在线监测')
    
    
    
    # 生成表格
    table = report.add_table(rows=4,  cols=6)
    
    #设定单元格字体和内容
    def set_cell_text(cell, text):
        p = cell.paragraphs[0]
        run = p.add_run()
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.add_text(text)
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = True
    
    
    set_cell_text(table.cell(0, 0), '产品名称')
    set_cell_text(table.cell(0, 1), args.product)  
    set_cell_text(table.cell(0, 2), '样品编号')
    set_cell_text(table.cell(0, 3), args.samplenumber)
    set_cell_text(table.cell(0, 4), '型号规格')
    set_cell_text(table.cell(0, 5), args.type)
    
    set_cell_text(table.cell(1, 0), '受检单位')
    set_cell_text(table.cell(1, 1), args.institution)
    set_cell_text(table.cell(1, 2), '设备名称')
    set_cell_text(table.cell(1, 3), '多通道免缠绕插回损测试仪（单模）MS08B')
    set_cell_text(table.cell(1, 4), '出厂编号')
    set_cell_text(table.cell(1, 5), '1538556')
    set_cell_text(table.cell(2, 0), '检验时间')
    set_cell_text(table.cell(3, 0), '检验人员')
    
    #合并一些表格单元
    time_cell = table.cell(2, 1).merge(table.cell(2,5))
    set_cell_text(time_cell, start_time_string + ' 至 ' + end_time_string)
    
    person_cell = table.cell(3, 1).merge(table.cell(3,5))
    p = person_cell.paragraphs[0]
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    
    
    report.add_paragraph(' ')
    
    for i in range(j * 2, j * 2 + 2):
        #插入图形的长宽
        report.add_picture(str(i) + '.jpg', width=Inches(1.5 * 4), height=Inches(1.5))
    
    report.save('报告'+str(j + 1) + '.docx')
    ##generate_report(args.path, args.product, args.samplenumber, args.institution)